#!/usr/bin/env python3
"""
Export publish markdown drafts to DOCX and PDF (structure-first draft style).

Default paths:
  - input:  /home/tj/eda_teaching_pack/publish/md
  - docx:   /home/tj/eda_teaching_pack/publish/docx
  - pdf:    /home/tj/eda_teaching_pack/publish/pdf
"""

from __future__ import annotations

import argparse
import html
import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.shared import Pt
from weasyprint import HTML


HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
ORDERED_RE = re.compile(r"^\d+\.\s+(.*)$")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Export markdown drafts to DOCX/PDF.")
    parser.add_argument("--input-dir", default="/home/tj/eda_teaching_pack/publish/md")
    parser.add_argument("--docx-dir", default="/home/tj/eda_teaching_pack/publish/docx")
    parser.add_argument("--pdf-dir", default="/home/tj/eda_teaching_pack/publish/pdf")
    parser.add_argument("--pattern", default="*.md", help="Glob pattern for markdown files.")
    return parser.parse_args()


def parse_table_lines(table_lines: list[str]) -> tuple[list[str], list[list[str]]]:
    rows: list[list[str]] = []
    for raw in table_lines:
        line = raw.strip()
        if not (line.startswith("|") and line.endswith("|")):
            continue
        cells = [c.strip() for c in line[1:-1].split("|")]
        rows.append(cells)

    if not rows:
        return [], []

    if len(rows) >= 2:
        delimiter = rows[1]
        if delimiter and all(set(cell.replace(":", "").strip()) <= {"-"} for cell in delimiter):
            header = rows[0]
            data = rows[2:]
            return header, data

    return rows[0], rows[1:]


def flush_table_docx(doc: Document, table_lines: list[str]) -> None:
    header, data_rows = parse_table_lines(table_lines)
    if not header:
        return
    n_cols = len(header)
    if n_cols == 0:
        return

    table = doc.add_table(rows=1, cols=n_cols)
    table.style = "Table Grid"
    for i, text in enumerate(header):
        table.cell(0, i).text = text

    for row in data_rows:
        r = table.add_row().cells
        for i in range(n_cols):
            r[i].text = row[i] if i < len(row) else ""

    doc.add_paragraph("")


def markdown_to_docx(md_text: str, out_path: Path) -> None:
    doc = Document()
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Malgun Gothic"
    normal_style.font.size = Pt(11)

    lines = md_text.splitlines()
    in_code = False
    table_lines: list[str] = []

    for idx, line in enumerate(lines + [""]):
        stripped = line.strip()

        if in_code:
            if stripped.startswith("```"):
                in_code = False
                doc.add_paragraph("")
            else:
                p = doc.add_paragraph(line.rstrip())
                for run in p.runs:
                    run.font.name = "Consolas"
                    run.font.size = Pt(9)
            continue

        if stripped.startswith("```"):
            in_code = True
            continue

        is_table_line = stripped.startswith("|") and stripped.endswith("|")
        if is_table_line:
            table_lines.append(line)
            continue

        if table_lines:
            flush_table_docx(doc, table_lines)
            table_lines = []

        if not stripped:
            if idx != len(lines):
                doc.add_paragraph("")
            continue

        m = HEADING_RE.match(stripped)
        if m:
            level = min(len(m.group(1)), 4)
            text = m.group(2).strip()
            doc.add_heading(text, level=level)
            continue

        m = ORDERED_RE.match(stripped)
        if m:
            doc.add_paragraph(m.group(1).strip(), style="List Number")
            continue

        if stripped.startswith("- "):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
            continue

        doc.add_paragraph(stripped)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def flush_table_html(parts: list[str], table_lines: list[str]) -> None:
    header, data_rows = parse_table_lines(table_lines)
    if not header:
        return
    parts.append("<table>")
    parts.append("<thead><tr>")
    for cell in header:
        parts.append(f"<th>{html.escape(cell)}</th>")
    parts.append("</tr></thead>")
    parts.append("<tbody>")
    for row in data_rows:
        parts.append("<tr>")
        for cell in row:
            parts.append(f"<td>{html.escape(cell)}</td>")
        parts.append("</tr>")
    parts.append("</tbody></table>")


def markdown_to_html(md_text: str, title: str) -> str:
    lines = md_text.splitlines()
    parts: list[str] = []
    in_code = False
    in_ul = False
    in_ol = False
    table_lines: list[str] = []

    def close_lists() -> None:
        nonlocal in_ul, in_ol
        if in_ul:
            parts.append("</ul>")
            in_ul = False
        if in_ol:
            parts.append("</ol>")
            in_ol = False

    def flush_table() -> None:
        nonlocal table_lines
        if table_lines:
            close_lists()
            flush_table_html(parts, table_lines)
            table_lines = []

    for line in lines:
        stripped = line.strip()

        if in_code:
            if stripped.startswith("```"):
                parts.append("</code></pre>")
                in_code = False
            else:
                parts.append(html.escape(line))
            continue

        if stripped.startswith("```"):
            flush_table()
            close_lists()
            parts.append("<pre><code>")
            in_code = True
            continue

        is_table_line = stripped.startswith("|") and stripped.endswith("|")
        if is_table_line:
            table_lines.append(line)
            continue
        flush_table()

        if not stripped:
            close_lists()
            continue

        m = HEADING_RE.match(stripped)
        if m:
            close_lists()
            level = min(len(m.group(1)), 4)
            parts.append(f"<h{level}>{html.escape(m.group(2).strip())}</h{level}>")
            continue

        m = ORDERED_RE.match(stripped)
        if m:
            if in_ul:
                parts.append("</ul>")
                in_ul = False
            if not in_ol:
                parts.append("<ol>")
                in_ol = True
            parts.append(f"<li>{html.escape(m.group(1).strip())}</li>")
            continue

        if stripped.startswith("- "):
            if in_ol:
                parts.append("</ol>")
                in_ol = False
            if not in_ul:
                parts.append("<ul>")
                in_ul = True
            parts.append(f"<li>{html.escape(stripped[2:].strip())}</li>")
            continue

        close_lists()
        parts.append(f"<p>{html.escape(stripped)}</p>")

    flush_table()
    close_lists()

    css = """
    @page { size: A4; margin: 18mm 14mm; }
    body {
      font-family: "Noto Sans CJK KR", "Malgun Gothic", "Apple SD Gothic Neo", sans-serif;
      font-size: 11pt;
      line-height: 1.55;
      color: #111;
    }
    h1 { font-size: 22pt; margin: 0 0 10pt 0; }
    h2 { font-size: 15pt; margin: 16pt 0 8pt 0; }
    h3 { font-size: 12.5pt; margin: 12pt 0 6pt 0; }
    h4 { font-size: 11.5pt; margin: 10pt 0 6pt 0; }
    p { margin: 0 0 6pt 0; }
    ul, ol { margin: 0 0 8pt 16pt; }
    li { margin: 0 0 4pt 0; }
    pre {
      background: #f5f5f5;
      border: 1px solid #e0e0e0;
      padding: 8pt;
      font-size: 9.5pt;
      white-space: pre-wrap;
    }
    table {
      border-collapse: collapse;
      width: 100%;
      margin: 8pt 0 10pt 0;
      font-size: 10pt;
    }
    th, td {
      border: 1px solid #d6d6d6;
      padding: 5pt 6pt;
      vertical-align: top;
      text-align: left;
    }
    th {
      background: #f1f3f5;
      font-weight: 700;
    }
    """

    return (
        "<!doctype html><html><head><meta charset='utf-8'>"
        f"<title>{html.escape(title)}</title><style>{css}</style></head><body>"
        + "".join(parts)
        + "</body></html>"
    )


def export_one(md_path: Path, docx_dir: Path, pdf_dir: Path) -> tuple[Path, Path]:
    stem = md_path.stem
    docx_path = docx_dir / f"{stem}.docx"
    pdf_path = pdf_dir / f"{stem}.pdf"

    md_text = md_path.read_text(encoding="utf-8")
    markdown_to_docx(md_text, docx_path)
    html_text = markdown_to_html(md_text, title=stem)
    HTML(string=html_text, base_url=str(md_path.parent)).write_pdf(str(pdf_path))
    return docx_path, pdf_path


def iter_markdown_files(input_dir: Path, pattern: str) -> Iterable[Path]:
    return sorted(p for p in input_dir.glob(pattern) if p.is_file())


def main() -> None:
    args = parse_args()
    input_dir = Path(args.input_dir)
    docx_dir = Path(args.docx_dir)
    pdf_dir = Path(args.pdf_dir)

    docx_dir.mkdir(parents=True, exist_ok=True)
    pdf_dir.mkdir(parents=True, exist_ok=True)

    files = list(iter_markdown_files(input_dir, args.pattern))
    if not files:
        raise SystemExit(f"No markdown files found in {input_dir} with pattern {args.pattern}")

    print(f"Exporting {len(files)} markdown files...")
    for md_path in files:
        docx_path, pdf_path = export_one(md_path, docx_dir, pdf_dir)
        print(f"- {md_path.name} -> {docx_path.name}, {pdf_path.name}")

    print("Done")


if __name__ == "__main__":
    main()
